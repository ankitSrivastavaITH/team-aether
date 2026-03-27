#!/usr/bin/env python3
"""Generate RVA Contract Lens Team Brief as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)
    return p

# ==========================================
# TITLE PAGE
# ==========================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('RVA Contract Lens', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.font.size = Pt(36)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Team Brief & Business Plan')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Hack for RVA 2026 | March 27–29\n').font.size = Pt(14)
details.add_run('VCU School of Business, Richmond, VA\n').font.size = Pt(12)
details.add_run('\n')
details.add_run('Track 1: A Thriving City Hall\n').font.size = Pt(12)
details.add_run('Targeting: Pillar Award + Moonshot (People\'s Choice)').font.size = Pt(12)

doc.add_page_break()

# ==========================================
# EXECUTIVE SUMMARY
# ==========================================
add_heading_styled('Executive Summary', 1)
doc.add_paragraph(
    'RVA Contract Lens is a dual-purpose civic technology tool that transforms how Richmond '
    'manages and communicates public spending. For City procurement staff, it replaces hours of '
    'manual CSV/PDF review with an intelligent dashboard that surfaces expiring contracts, '
    'extracts key terms from procurement documents, and answers natural language questions about '
    'contract data. For Richmond residents, it provides an unprecedented window into how $6.1 '
    'billion in public contracts are awarded, to whom, and when they expire.'
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Target Awards:')
run.bold = True
add_bullet('Staff-facing procurement risk dashboard', 'Pillar Award (Track 1):')
add_bullet('"Where do your tax dollars go?"', 'Moonshot Award (People\'s Choice):')

# ==========================================
# THE PROBLEM
# ==========================================
add_heading_styled('The Problem', 1)

add_heading_styled('For City Staff', 2)
doc.add_paragraph(
    'Richmond\'s procurement team manages 1,365+ contracts across 15 departments totaling $6.1B. '
    'To make informed purchasing decisions, staff must:'
)
add_bullet('Download a CSV from the City\'s open data portal')
add_bullet('Open it in Excel, manually sort and filter by dates')
add_bullet('Cross-reference with SAM.gov (federal), eVA (state), and VITA (state IT) portals')
add_bullet('Open 40-page procurement PDFs to find renewal windows, pricing terms, and conditions')
add_bullet('Repeat for every contract review cycle')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Scale of the pain today:')
run.bold = True
add_bullet('22 contracts expiring in the next 30 days')
add_bullet('54 contracts expiring in the next 60 days')
add_bullet('75 contracts expiring in the next 90 days')
add_bullet('No automated alerting for any of these')

add_heading_styled('For Residents', 2)
doc.add_paragraph(
    'Richmond\'s $6.1B in public contracts is technically public data, but practically invisible. '
    'The raw CSV on the open data portal is unusable for anyone without data analysis skills. '
    'Residents have no way to:'
)
add_bullet('See which vendors receive the most public money')
add_bullet('Track spending trends by department over time')
add_bullet('Understand what contracts are active in their city')
add_bullet('Hold the City accountable for procurement decisions')

# ==========================================
# THE SOLUTION
# ==========================================
add_heading_styled('The Solution', 1)
doc.add_paragraph('One app. Two views. Two awards.')

add_heading_styled('Staff View — Procurement Risk Dashboard', 2)
add_table(
    ['Feature', 'Description'],
    [
        ['Expiry Tracker', 'Filterable table, color-coded: Red (≤30d), Yellow (31-60d), Green (60d+)'],
        ['PDF Extractor', 'Upload contract PDF → AI extracts expiration, renewal, pricing, conditions'],
        ['NL Query Bar', '"Show me all DPW contracts over $500K expiring this quarter" → table + chart'],
        ['Risk Alerts', 'Auto-generated: contracts expiring soon, high-value renewals'],
    ]
)

add_heading_styled('Public View — "Where Do Your Tax Dollars Go?"', 2)
add_table(
    ['Feature', 'Description'],
    [
        ['Spending Explorer', '$6.1B visualized by department, vendor, year — interactive charts'],
        ['Vendor Explorer', 'Click any vendor → full contract history, cross-reference federal data'],
        ['Contract Search', 'Plain-language summaries, full provenance, links to official sources'],
        ['Trust Labels', 'Every page: "Exploratory tool — not official City financial reporting"'],
    ]
)

# ==========================================
# TECH STACK
# ==========================================
add_heading_styled('Technical Architecture', 1)
add_table(
    ['Layer', 'Technology', 'Why'],
    [
        ['Frontend', 'Next.js + shadcn/ui + Recharts + TanStack Table', 'Team expertise, fast to build'],
        ['Backend', 'FastAPI + DuckDB (in-process)', 'Zero setup, native CSV loading, analytical SQL'],
        ['AI', 'Groq (free 500k tokens/day)', 'Free tier, fastest inference for live demos'],
        ['Data', 'City Contracts CSV + SAM.gov API + eVA CSV', 'All public, validated'],
        ['Deploy', 'Vercel (frontend) + local FastAPI', 'Free hosting, instant deploys'],
    ]
)

# ==========================================
# DATA FOUNDATION
# ==========================================
add_heading_styled('Data Foundation (Validated)', 1)
add_table(
    ['Source', 'Records', 'Status'],
    [
        ['City Contracts (Socrata CSV)', '1,365', 'Downloaded & validated'],
        ['SAM.gov', 'TBD', 'API key needed day-1 (free)'],
        ['eVA (Virginia)', 'TBD', 'CSV download day-1'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Key Metrics From Real Data:')
run.bold = True
add_bullet('$6.1B total contract value')
add_bullet('15 departments represented')
add_bullet('22 contracts expiring in 30 days')
add_bullet('Top vendor: Carahsoft Technology (24 contracts)')
add_bullet('Largest single contract: $810.6M')
add_bullet('Date range: 2011 to 2031+')
add_bullet('100% of dates parseable — no data quality issues')

p = doc.add_paragraph()
run = p.add_run('⚠ Do NOT use the Socrata API — known bug returns only 8 of 9 columns. Use the CSV download.')
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True

# ==========================================
# BUILD TIMELINE
# ==========================================
add_heading_styled('Build Timeline', 1)

add_heading_styled('Friday 3/27 — Foundation', 2)
add_table(
    ['Task', 'Hours'],
    [
        ['Scaffold Next.js + FastAPI project', '1'],
        ['Download CSVs, register SAM.gov API key', '0.5'],
        ['Write data ingestion script (CSV → DuckDB)', '2'],
        ['Source 3-5 sample procurement PDFs', '1'],
        ['Deploy frontend shell to Vercel', '0.5'],
        ['Set up project repo on GitHub', '0.5'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('Friday exit criteria: ')
run.bold = True
p.add_run('DuckDB loaded with all contract data, frontend shell live on Vercel.')

add_heading_styled('Saturday 3/28 — Build Day', 2)
add_table(
    ['Time', 'Task', 'Priority'],
    [
        ['8-10 AM', 'Staff dashboard: expiry table + filters', 'P0'],
        ['10-12 PM', 'Risk alerts panel + contract detail view', 'P0'],
        ['12-2 PM', 'PDF extraction endpoint + upload UI', 'P1'],
        ['2-4 PM', 'NL-to-SQL query bar', 'P1'],
        ['4-6 PM', 'Public transparency view: spending charts', 'P0'],
        ['6-8 PM', 'Vendor explorer + contract search', 'P1'],
        ['8-10 PM', 'Polish, bug fixes, demo prep', 'P0'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('P0 = must have for demo. P1 = strong nice-to-have.')
run.bold = True

add_heading_styled('Sunday 3/29 — Demo Day', 2)
add_table(
    ['Time', 'Task'],
    [
        ['Morning', 'Polish demo flow, pre-cache queries'],
        ['11 AM', 'Test on projector/screen resolution'],
        ['12 PM', 'Practice pitch (3 min, timed)'],
        ['1:30 PM', 'Finals at VCU'],
    ]
)

# ==========================================
# PRIORITY STACK
# ==========================================
add_heading_styled('Priority Stack (If We Run Out of Time)', 1)
add_table(
    ['Priority', 'Feature', 'Why'],
    [
        ['P0 - Must Ship', 'Expiry tracker table with filters', 'Core value proposition for staff'],
        ['P0 - Must Ship', 'Spending charts in public view', 'Moonshot hook for audience vote'],
        ['P0 - Must Ship', 'Real data from City Contracts CSV', 'Credibility with judges'],
        ['P1 - Should Ship', 'PDF extraction', 'Strong Innovation score'],
        ['P1 - Should Ship', 'NL-to-SQL query bar', 'Wow factor for demo'],
        ['P1 - Should Ship', 'Vendor cross-reference (SAM.gov)', 'Multi-source story'],
        ['P2 - Nice to Have', 'eVA state data integration', 'Broader data coverage'],
        ['P2 - Nice to Have', 'AI plain-language contract summaries', 'Public view polish'],
        ['P2 - Nice to Have', 'Department-level risk aggregation', 'Staff view depth'],
    ]
)

# ==========================================
# PITCH SCRIPTS
# ==========================================
add_heading_styled('Pitch Structure (3 Minutes)', 1)

add_heading_styled('Pillar Award — For Judges', 2)
add_table(
    ['Time', 'Content'],
    [
        ['0:00-0:30', 'THE PAIN: "Procurement officers download CSVs, open Excel, scan PDFs. Hours per review cycle."'],
        ['0:30-2:00', 'STAFF DEMO: Filter expiring contracts → click one → upload PDF → AI extracts → NL query'],
        ['2:00-2:30', 'CREDIBILITY: "Real City data. AI labeled advisory. Staff make the final call."'],
        ['2:30-3:00', 'CLOSE: "RVA Contract Lens: less time in spreadsheets, fewer missed renewals."'],
    ]
)

add_heading_styled('Moonshot Award — For Audience Vote', 2)
add_table(
    ['Time', 'Content'],
    [
        ['0:00-0:30', 'THE HOOK: "Richmond spends $6.1 billion in public contracts. Can you name a single one?"'],
        ['0:30-2:00', 'PUBLIC DEMO: Spending by department → click vendor → history → federal cross-reference'],
        ['2:00-2:30', 'EMOTIONAL CLOSE: "This data was always public. We just made it visible."'],
        ['2:30-3:00', 'CTA: "Vote for RVA Contract Lens."'],
    ]
)

# ==========================================
# MOONSHOT STRATEGY
# ==========================================
add_heading_styled('Moonshot Strategy', 1)
doc.add_paragraph(
    'The People\'s Choice / Moonshot Award is won by audience vote. The audience is hackathon '
    'attendees — technologists, designers, community members, mentors, and city officials.'
)
p = doc.add_paragraph()
run = p.add_run('Why the transparency view wins votes:')
run.bold = True
add_bullet('Everyone pays taxes. Everyone wants to know where the money goes.', 'Universal appeal:')
add_bullet('$6.1 BILLION displayed visually is immediately compelling.', 'Visceral impact:')
add_bullet('Attendees can explore their city\'s spending in real-time during the demo.', 'Interactive discovery:')
add_bullet('"Your city spends $810M on its largest contract. Do you know what it\'s for?"', 'Emotional resonance:')
add_bullet('Richmond\'s first civic hackathon producing Richmond\'s first spending transparency tool.', 'Civic pride:')

# ==========================================
# RULES OF ENGAGEMENT
# ==========================================
add_heading_styled('Rules of Engagement', 1)

add_heading_styled('We DO', 2)
add_bullet('Use real, verified public data from Richmond\'s open data portal')
add_bullet('Label every AI output as "AI-assisted — verify against original"')
add_bullet('Link to official sources for every data point')
add_bullet('Say "Exploratory tool — not official City financial reporting" on every page')
add_bullet('Support staff judgment — never replace it')

add_heading_styled('We DO NOT', 2)
add_bullet('Integrate with City systems (RVA311, EnerGov, Oracle RAPIDS)')
add_bullet('Make compliance or legal determinations')
add_bullet('Make contract award or eligibility decisions')
add_bullet('Claim the tool is authoritative or deployment-ready')
add_bullet('Use synthetic/fake data when real data exists')

# ==========================================
# JUDGING RUBRIC
# ==========================================
add_heading_styled('Judging Rubric (Max 105 Points)', 1)
add_table(
    ['Category', 'Weight', 'Our Angle'],
    [
        ['Impact', '5x', 'Direct hit on Problem 2 + Blue Sky'],
        ['User Value', '4x (tiebreaker)', 'Staff save hours; residents see spending'],
        ['Feasibility', '3x', 'Public data, standard stack, no dependencies'],
        ['Innovation', '3x', 'Dual view, NL queries, AI extraction'],
        ['Execution', '3x', 'Real data, coherent flow, two complete views'],
        ['Equity', '3x', 'Public transparency for all residents'],
    ]
)

# ==========================================
# POST-HACKATHON
# ==========================================
add_heading_styled('Post-Hackathon Roadmap', 1)
add_table(
    ['Phase', 'Timeline', 'Goals'],
    [
        ['Pilot', '0-3 months', 'Partner with City procurement, real-time data refresh, user testing'],
        ['Expand', '3-6 months', 'SAM.gov debarment checks, renewal reminders, mobile public view'],
        ['Platform', '6-12 months', 'Multi-city expansion (Virginia), API for civic tools'],
        ['Vision', '12+ months', 'Civic innovation lab — sustained procurement transparency'],
    ]
)

# ==========================================
# KEY LINKS
# ==========================================
add_heading_styled('Key Links', 1)
add_table(
    ['Resource', 'URL'],
    [
        ['Pillar repo', 'github.com/hack4rva/pillar-thriving-city-hall'],
        ['City Contracts CSV', 'data.richmondgov.com/api/views/xqn7-jvv2/rows.csv?accessType=DOWNLOAD'],
        ['SAM.gov API', 'api.sam.gov'],
        ['eVA data', 'data.virginia.gov'],
        ['Hackathon site', 'rvahacks.org'],
    ]
)

# ==========================================
# FOOTER
# ==========================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"This data was always public. We just made it visible."')
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Built for Hack for RVA 2026 — Track 1: A Thriving City Hall').font.size = Pt(10)

# Save
output_path = os.path.expanduser('/Users/ithena/Documents/CodeSpace/hackrva/RVA_Contract_Lens_Team_Brief.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
